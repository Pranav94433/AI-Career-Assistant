import streamlit as st
try:
    from groq import Groq
except ImportError:
    Groq = None
import io
import json
import os
import re
import base64
import importlib
import smtplib
import uuid
from datetime import datetime
from email.message import EmailMessage
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# ============================================================
# CONFIGURATION
# ============================================================

st.set_page_config(
    page_title="AI Career Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&family=Space+Grotesk:wght@500;600;700&display=swap');

:root {
    --ink: #f5f5f5;
    --muted: #b5b5b5;
    --mint: #263c37;
    --teal: #167d72;
    --coral: #f47b61;
    --paper: #050505;
}

html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
    color: var(--ink);
}

[data-testid="stAppViewContainer"] {
    background: var(--paper);
}

[data-testid="stHeader"] {
    background: transparent;
}

/* Hide Streamlit and deployment branding controls. */
#MainMenu,
footer,
[data-testid="stDeployButton"],
[data-testid="stStatusWidget"],
[data-testid="stDecoration"],
.stAppDeployButton,
.stStatusWidget,
[class*="stAppDeployButton"],
[class*="stStatusWidget"] {
    visibility: hidden;
    display: none;
}

[data-testid="stAppViewContainer"] * {
    color: var(--ink);
}

[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #151515 0%, #0b0b0b 100%);
    color: var(--ink);
}

[data-testid="stFileUploader"] {
    background: #1b1b1b;
    border: 2px solid #3f9f7f;
    border-radius: 8px;
    box-shadow: 0 5px 16px rgba(22, 125, 114, 0.16);
    padding: 0.8rem;
    width: 100%;
}

[data-testid="stMainBlockContainer"] {
    width: min(100%, 52rem);
    max-width: 52rem;
    padding-top: 2rem;
    padding-bottom: 8rem;
}

[data-testid="stHorizontalBlock"] {
    width: 100%;
}

[data-testid="stFileUploader"] section,
[data-testid="stFileUploaderDropzone"] {
    background: #111111;
    border: 2px dashed #5eb997;
    border-radius: 6px;
    min-height: 6rem;
    width: 100%;
}

[data-testid="stFileUploader"] small,
[data-testid="stFileUploader"] span,
[data-testid="stFileUploader"] button {
    color: #f5f5f5 !important;
}

[data-testid="stFileUploader"] button {
    background: var(--teal) !important;
    border: 0 !important;
    border-radius: 6px !important;
    color: #ffffff !important;
    font-weight: 700;
    min-height: 2.75rem;
    min-width: 8rem;
}

[data-testid="stExpander"] {
    border: 1px solid #3f9f7f;
    border-radius: 10px;
    background: #101817;
}

[data-testid="stExpander"] summary {
    color: #8de0c2 !important;
    font-weight: 700;
}

[data-testid="stChatMessage"] {
    background: transparent;
    border: 0;
    color: var(--ink) !important;
}

[data-testid="stChatMessage"][aria-label="assistant message"] {
    background: #171717;
    border: 1px solid #292929;
    border-radius: 8px;
    padding: 0.75rem 1rem;
}

[data-testid="stChatMessage"][aria-label="assistant message"] * {
    color: #ffffff !important;
    font-size: 0.94rem;
    line-height: 1.45;
}

[data-testid="stChatMessage"][aria-label="assistant message"] code {
    background: #292929;
}

[data-testid="stChatMessage"] p,
[data-testid="stChatMessage"] li,
[data-testid="stChatMessage"] strong,
[data-testid="stChatMessage"] code {
    color: var(--ink) !important;
}

[data-testid="stChatInput"] {
    background: #161616;
    border: 1px solid #424242;
    border-radius: 18px;
    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
    padding: 0.35rem 0.55rem;
}

[data-testid="stChatInput"] textarea {
    background: #161616 !important;
    border: 0 !important;
    border-radius: 14px !important;
    color: #ffffff !important;
    caret-color: var(--teal);
    min-height: 4.5rem !important;
    padding: 0.8rem 0.85rem !important;
    box-shadow: none;
    transition: background 160ms ease, border-color 160ms ease, box-shadow 160ms ease;
}

[data-testid="stChatInput"] textarea:focus {
    background: #1d1d1d !important;
    border-color: transparent !important;
    box-shadow: none;
}

[data-testid="stChatInputSubmitButton"] button {
    background: #3f9f7f !important;
    border-radius: 50% !important;
    color: #ffffff !important;
}

[data-testid="stChatInput"] textarea::placeholder {
    color: var(--muted) !important;
    opacity: 1;
}

h1, h2, h3 {
    font-family: 'Space Grotesk', sans-serif;
    letter-spacing: 0;
}

.welcome {
    padding: 1.5rem 0 0.7rem;
    text-align: center;
}

.welcome-kicker {
    color: var(--teal);
    font-size: 0.8rem;
    font-weight: 700;
    letter-spacing: 0.08em;
    text-transform: uppercase;
}

.welcome h2 {
    color: var(--ink);
    font-size: clamp(1.8rem, 4vw, 3.1rem);
    line-height: 1.05;
    margin: 0.35rem 0 0.65rem;
}

.welcome p {
    color: var(--muted);
    font-size: 1.05rem;
    max-width: 42rem;
    margin: 0 auto;
}

div.stButton > button {
    border: 1px solid #3b514a;
    border-radius: 8px;
    color: #f5f5f5;
    background: #1b2b27;
    min-height: 3.2rem;
    text-align: left;
    transition: border-color 160ms ease, transform 160ms ease, background 160ms ease;
}

@media (max-width: 900px) {
    [data-testid="stMainBlockContainer"] {
        padding: 1.25rem 1rem 7rem;
    }

    .welcome h2 {
        font-size: 2.2rem;
    }
}

@media (max-width: 600px) {
    [data-testid="stMainBlockContainer"] {
        padding: 0.75rem 0.7rem 6rem;
    }

    h1 {
        font-size: 1.65rem !important;
    }

    .welcome {
        padding-top: 0.75rem;
    }

    .welcome h2 {
        font-size: 1.8rem;
    }

    .welcome p {
        font-size: 0.95rem;
    }

    [data-testid="stHorizontalBlock"] {
        flex-wrap: wrap;
        gap: 0.5rem;
    }

    [data-testid="stHorizontalBlock"] > [data-testid="column"] {
        flex: 1 1 100% !important;
        min-width: 100% !important;
    }

    [data-testid="stChatMessage"] {
        padding-left: 0.45rem;
        padding-right: 0.45rem;
    }

    [data-testid="stChatInput"] textarea {
        min-height: 3.8rem !important;
        padding: 0.7rem !important;
    }

    [data-testid="stFileUploader"] {
        padding: 0.55rem;
    }
}

div.stButton > button:hover {
    border-color: var(--teal);
    background: var(--mint);
    transform: translateY(-2px);
}

[data-testid="stChatMessage"] {
    border-radius: 10px;
}
</style>
""", unsafe_allow_html=True)

PREFERRED_MODELS = [
    "openai/gpt-oss-20b",
    "llama-4-scout-17b-16e-instruct",
    "qwen/qwen3-32b",
]

KNOWLEDGE_FILE = "universal_ai_career_assistant_knowledge_base.txt"
REVIEWS_FILE = "career_assistant_reviews.json"
FEEDBACK_RECIPIENT = "pranavkumar86530@gmail.com"


# ============================================================
# LOAD KNOWLEDGE BASE
# ============================================================

@st.cache_resource
def load_knowledge_base():

    file_path = Path(KNOWLEDGE_FILE)

    if not file_path.exists():
        return [], None, None

    text = file_path.read_text(
        encoding="utf-8",
        errors="ignore"
    )

    chunks = []

    chunk_size = 1000
    overlap = 200

    start = 0

    while start < len(text):

        end = start + chunk_size

        chunk = text[start:end]

        chunks.append(chunk)

        start += chunk_size - overlap

    vectorizer = TfidfVectorizer(
        stop_words="english"
    )

    vectors = vectorizer.fit_transform(chunks)

    return chunks, vectorizer, vectors


# ============================================================
# SEARCH RELEVANT KNOWLEDGE
# ============================================================

def search_knowledge(question, chunks, vectorizer, vectors):

    if not chunks:
        return ""

    question_vector = vectorizer.transform(
        [question]
    )

    similarities = cosine_similarity(
        question_vector,
        vectors
    ).flatten()

    top_results = similarities.argsort()[-2:][::-1]

    relevant_chunks = []

    for index in top_results:

        if similarities[index] > 0:

            relevant_chunks.append(
                chunks[index]
            )

    return "\n\n".join(relevant_chunks)


@st.cache_data(ttl=900, show_spinner=False)
def search_web(question, resume_text=""):

    try:
        search_client = importlib.import_module("ddgs").DDGS
    except ImportError:
        return [], "The live web search package is not installed."

    resume_hint = " ".join(resume_text.split())[:1800]
    resume_hint = re.sub(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", "", resume_hint)
    resume_hint = re.sub(r"(?:\+?\d[\d ()-]{7,}\d)", "", resume_hint)
    current_year = datetime.now().year
    if resume_hint:
        query = (
            f"{question} current {current_year} job openings companies hiring roles "
            f"that match these skills: {resume_hint}"
        )
    else:
        query = (
            f"{question} best companies employers industry leaders "
            f"current {current_year} official company information"
        )

    try:
        results = list(search_client().text(query, max_results=5))
    except Exception as error:
        return [], f"Live web search was unavailable: {error}"

    usable_results = [
        {
            "title": result.get("title", "Untitled result"),
            "url": result.get("href", ""),
            "snippet": result.get("body", ""),
        }
        for result in results
        if result.get("href")
    ]
    web_context = "\n\n".join(
        f"Title: {result['title']}\nURL: {result['url']}\nSnippet: {result['snippet']}"
        for result in usable_results
    )
    return usable_results, web_context


# ============================================================
# AI RESPONSE
# ============================================================

def get_ai_response(question, knowledge_context, conversation_history):

    memory_context = "\n".join(
        f"{message['role'].title()}: {message['content']}"
        for message in conversation_history[-8:]
    )

    system_prompt = f"""
You are a warm, practical AI career coach. Make the conversation feel personal,
useful, and encouraging without being vague or overconfident. Answer the exact
question first, then give clear next steps. When an uploaded resume is present,
refer to its actual details and clearly say when information is missing. Do not
guarantee jobs or salaries and do not invent qualifications.

Context:

{knowledge_context}

Recent conversation memory:
{memory_context or "No previous conversation yet."}

Answer in a friendly structure with a direct answer, a brief explanation,
practical steps or examples, and a thoughtful follow-up question when it would
help the user continue. Give enough detail to be genuinely useful, usually
around 6-10 bullet points or short paragraphs. Never stop after only one sentence.

When live web results are provided, use them for current companies, roles,
salary trends, and hiring information. Clearly label time-sensitive claims,
tell the user to verify that a role is still open, and do not claim that a
company is hiring based only on a search snippet. Never invent job listings or
resume matches.
"""

    groq_client = get_groq_client()
    model_name = get_available_model(groq_client)

    return groq_client.chat.completions.create(

        model=model_name,

        stream=True,

        temperature=0.2,
        max_tokens=1200,

        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": question
            }
        ]

    )


def generate_ats_resume(source_resume, additional_details):

    source_text = source_resume.strip() or "No existing resume was uploaded."
    details_text = additional_details.strip() or "No additional details were provided."
    prompt = f"""
Create a polished, ATS-friendly resume from the candidate information below.
Use only facts explicitly provided. Never invent employers, dates, degrees,
certifications, job titles, metrics, skills, or contact information. If an
important detail is missing, omit it or use a clear placeholder such as
[Phone Number] rather than guessing.

Return only the resume text, with no markdown fences, commentary, or explanation.
Use this order when the information exists:
1. Full name and contact information
2. Professional summary tailored to the candidate's target role
3. Core skills using standard searchable keywords
4. Professional experience with concise achievement-focused bullet points
5. Projects
6. Education
7. Certifications and additional information

Keep formatting plain and ATS-compatible: simple headings, no tables, columns,
icons, graphics, emojis, or decorative characters. Strengthen vague bullets
only by rewriting the supplied facts; do not add unsupported accomplishments.

Existing resume:
{source_text[:16000]}

Additional candidate details and target role:
{details_text[:8000]}
"""

    groq_client = get_groq_client()
    model_name = get_available_model(groq_client)
    response = groq_client.chat.completions.create(
        model=model_name,
        temperature=0.1,
        max_tokens=2500,
        messages=[
            {
                "role": "system",
                "content": "You are an expert ATS resume writer who preserves factual accuracy.",
            },
            {"role": "user", "content": prompt},
        ],
    )
    return response.choices[0].message.content.strip()


def create_resume_docx(resume_text):

    document = Document()
    for line in resume_text.splitlines():
        stripped_line = line.strip()
        if not stripped_line:
            document.add_paragraph()
        elif stripped_line.isupper() or stripped_line.endswith(":"):
            document.add_heading(stripped_line.rstrip(":"), level=2)
        else:
            document.add_paragraph(stripped_line)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


def extract_resume_text(uploaded_file):

    file_type = Path(uploaded_file.name).suffix.lower()
    file_bytes = uploaded_file.getvalue()

    if file_type == ".pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        if reader.is_encrypted:
            raise ValueError("This PDF is password-protected. Remove the password and upload it again.")
        resume_text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if not resume_text:
            raise ValueError("No selectable text was found. This may be a scanned PDF; upload a text-based PDF or TXT/DOCX file.")
        return resume_text

    if file_type == ".docx":
        document = Document(io.BytesIO(file_bytes))
        paragraph_text = [paragraph.text for paragraph in document.paragraphs]
        table_text = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        resume_text = "\n".join(paragraph_text + table_text).strip()
        if not resume_text:
            raise ValueError("The DOCX file contains no readable text. Please check the document and upload it again.")
        return resume_text

    if file_type == ".txt":
        resume_text = file_bytes.decode("utf-8", errors="ignore").strip()
        if not resume_text:
            raise ValueError("The TXT file is empty. Add resume text and upload it again.")
        return resume_text

    if file_type in {".jpg", ".jpeg", ".png", ".webp"}:
        groq_client = get_groq_client()
        available_models = {
            model.id for model in groq_client.models.list().data
        }
        vision_model = "meta-llama/llama-4-scout-17b-16e-instruct"
        if vision_model not in available_models:
            raise ValueError("Image resumes require the Llama 4 Scout model. Upload a PDF, DOCX, or TXT resume instead.")

        mime_type = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".webp": "image/webp",
        }[file_type]
        image_data = base64.b64encode(file_bytes).decode("ascii")
        response = groq_client.chat.completions.create(
            model=vision_model,
            temperature=0,
            max_tokens=2000,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Extract all readable text from this resume image. Preserve headings, dates, skills, education, experience, and project details. Return only the extracted text.",
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{image_data}",
                            },
                        },
                    ],
                }
            ],
        )
        resume_text = response.choices[0].message.content.strip()
        if not resume_text:
            raise ValueError("No readable text was found in the resume image. Upload a clearer image or a PDF/DOCX/TXT file.")
        return resume_text

    raise ValueError("Unsupported file type. Please upload a PDF, DOCX, TXT, JPG, JPEG, PNG, or WEBP file.")


def process_uploaded_resume(uploaded_resume):

    if uploaded_resume is None:
        return

    resume_id = f"{uploaded_resume.name}:{uploaded_resume.size}"
    if st.session_state.get("resume_id") == resume_id:
        return

    try:
        resume_text = extract_resume_text(uploaded_resume)
        st.session_state.resume_text = resume_text
        st.session_state.resume_id = resume_id
        st.session_state.resume_name = uploaded_resume.name
        st.session_state.resume_error = ""
    except Exception as error:
        st.session_state.resume_text = ""
        st.session_state.resume_id = resume_id
        st.session_state.resume_name = uploaded_resume.name
        st.session_state.resume_error = f"{type(error).__name__}: {error}"


def get_response_text(response_chunk):

    if isinstance(response_chunk, dict):
        return response_chunk.get("message", {}).get("content", "")

    choices = getattr(response_chunk, "choices", [])
    if choices:
        delta = getattr(choices[0], "delta", None)
        return getattr(delta, "content", "") or ""

    return ""


def get_setting(name, default=""):

    if name in os.environ:
        return os.environ[name]

    try:
        return st.secrets.get(name, default)
    except Exception:
        return default


@st.cache_resource
def get_groq_client():

    groq_api_key = get_setting("GROQ_API_KEY")
    if not groq_api_key:
        raise ValueError("GROQ_API_KEY is not configured")
    return Groq(api_key=groq_api_key)


def get_available_model(groq_client):

    available_models = {
        model.id for model in groq_client.models.list().data
    }

    for model_name in PREFERRED_MODELS:
        if model_name in available_models:
            return model_name

    raise ValueError(
        "No supported chat model is available for this Groq API key. "
        "Check the Models page in the Groq console."
    )


def send_feedback_email(rating, feedback):

    smtp_user = get_setting("SMTP_USER")
    smtp_password = get_setting("SMTP_PASSWORD")
    smtp_host = get_setting("SMTP_HOST", "smtp.gmail.com")
    smtp_port = int(get_setting("SMTP_PORT", "465"))

    if not smtp_user or not smtp_password:
        return False, "Email is not configured yet. Add SMTP_USER and SMTP_PASSWORD to send notifications."

    message = EmailMessage()
    message["Subject"] = f"AI Career Assistant feedback: {rating}/5"
    message["From"] = smtp_user
    message["To"] = FEEDBACK_RECIPIENT
    message.set_content(
        f"New chatbot feedback received.\n\n"
        f"Rating: {rating}/5\n"
        f"Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        f"Feedback:\n{feedback or '(No written feedback provided)'}\n"
    )

    try:
        with smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=15) as smtp:
            smtp.login(smtp_user, smtp_password)
            smtp.send_message(message)
    except Exception as error:
        return False, f"Email could not be sent: {error}"

    return True, "Feedback sent successfully. Thank you!"


def load_reviews():

    reviews_path = Path(REVIEWS_FILE)
    if not reviews_path.exists():
        return []

    try:
        reviews = json.loads(reviews_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []

    return reviews if isinstance(reviews, list) else []


def save_review(name, rating, feedback, owner_token):

    reviews = load_reviews()
    review_id = uuid.uuid4().hex
    reviews.insert(
        0,
        {
            "id": review_id,
            "owner_token": owner_token,
            "name": name.strip()[:80],
            "rating": int(rating),
            "feedback": feedback.strip()[:1000],
            "created_at": datetime.now().strftime("%Y-%m-%d"),
        },
    )
    Path(REVIEWS_FILE).write_text(
        json.dumps(reviews[:50], ensure_ascii=True, indent=2),
        encoding="utf-8",
    )


def delete_review(review_id, owner_token):

    reviews = load_reviews()
    matching_review = next(
        (
            review
            for review in reviews
            if review.get("id") == review_id
            and review.get("owner_token") == owner_token
        ),
        None,
    )
    if matching_review is None:
        return False

    remaining_reviews = [review for review in reviews if review is not matching_review]
    Path(REVIEWS_FILE).write_text(
        json.dumps(remaining_reviews, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )
    return True

# ============================================================
# SIDEBAR
# ============================================================

if "review_owner_token" not in st.session_state:
    st.session_state.review_owner_token = uuid.uuid4().hex

with st.sidebar:

    st.title("🤖 AI Career Assistant")

    st.write(
        "Your personal AI assistant for career guidance."
    )

    st.divider()

    st.divider()

    st.subheader("What can I help with?")

    live_search_enabled = st.checkbox(
        "Use current web results",
        value=True,
        help="Searches the public web for current companies, roles, and market information.",
    )

    topic_questions = {
        "Career guidance": "Give me personalized career guidance. Explain how I can choose a career direction based on my interests, education, experience, and current skills. Include practical next steps.",
        "Resume improvement": "Review my resume if I uploaded one and give detailed, actionable improvements for its content, structure, achievements, keywords, and formatting.",
        "Interview preparation": "Help me prepare for a job interview with common questions, answer frameworks, sample answers, and a practice plan.",
        "Career roadmaps": "Create a detailed 30-60-90 day career roadmap with skills to learn, projects to complete, and measurable goals.",
        "Skill recommendations": "Recommend the most valuable skills I should learn next for my target career. Explain why each skill matters and how to practice it.",
        "Job search guidance": "Give me a practical job-search strategy including suitable job boards, networking, applications, tracking, and ways to improve my chances.",
        "LinkedIn advice": "Explain how to improve my LinkedIn profile, headline, About section, skills, projects, networking, and recruiter visibility.",
        "Career changes": "Help me plan a realistic career change. Compare my transferable skills, identify gaps, suggest target roles, and create a step-by-step transition plan.",
    }

    for topic, topic_question in topic_questions.items():
        if st.button(topic, key=f"topic_{topic}", use_container_width=True):
            st.session_state.pending_question = topic_question
            st.rerun()

    st.divider()

    with st.expander("Rate this chatbot"):
        st.caption("Your feedback helps improve the career assistant.")

        with st.form("feedback_form", clear_on_submit=True):
            feedback_name = st.text_input(
                "Your name",
                placeholder="Enter your name",
                max_chars=80,
            )
            feedback_rating = st.radio(
                "How helpful was this chatbot?",
                options=[1, 2, 3, 4, 5],
                format_func=lambda value: "★" * value + "☆" * (5 - value),
                horizontal=True,
            )
            feedback_text = st.text_area(
                "Share your feedback (optional)",
                placeholder="What worked well, or what should be better?",
                max_chars=1000,
            )
            feedback_submitted = st.form_submit_button(
                "Send feedback",
                use_container_width=True,
            )

        if feedback_submitted:
            if not feedback_name.strip():
                st.warning("Please enter your name before sending a review.")
            else:
                save_review(
                    feedback_name,
                    feedback_rating,
                    feedback_text,
                    st.session_state.review_owner_token,
                )
                feedback_sent, feedback_message = send_feedback_email(
                    feedback_rating,
                    f"Name: {feedback_name.strip()}\n\n{feedback_text.strip()}",
                )
                if feedback_sent:
                    st.success("Your review was added. " + feedback_message)
                else:
                    st.success("Your review was added and is visible below.")

    with st.expander("See what others reviewed"):
        reviews = load_reviews()
        if not reviews:
            st.caption("No reviews yet. Be the first to share your experience.")
        else:
            for review in reviews:
                review_rating = max(1, min(5, int(review.get("rating", 0))))
                reviewer_name = review.get("name", "Anonymous")
                review_date = review.get("created_at", "")
                st.write(
                    f"{reviewer_name}  "
                    f"{'★' * review_rating}{'☆' * (5 - review_rating)}"
                    f"  {review_date}"
                )
                review_text = review.get("feedback", "")
                if review_text:
                    st.caption(review_text)
                if review.get("owner_token") == st.session_state.review_owner_token:
                    if st.button(
                        "Delete my review",
                        key=f"delete_review_{review.get('id', review_date)}",
                    ):
                        if delete_review(
                            review.get("id"),
                            st.session_state.review_owner_token,
                        ):
                            st.rerun()
                st.divider()

    if st.session_state.get("messages"):
        st.caption(f"🧠 Memory active: {len(st.session_state.messages)} messages")

        if st.button("Clear memory", key="clear_memory", use_container_width=True):
            st.session_state.messages = []
            st.session_state.pop("pending_question", None)
            st.rerun()

    st.divider()

    if st.button("🗑️ Clear Chat"):

        st.session_state.messages = []

        st.rerun()


# ============================================================
# MAIN PAGE
# ============================================================

st.title("🤖 AI Career Assistant Chatbot")

st.write(
    "Ask me about careers, skills, resumes, "
    "interviews, jobs, and professional development."
)
st.subheader("Upload your resume")
st.caption("Choose a resume from Files or your device. Supported formats: PDF, DOCX, TXT, JPG, JPEG, PNG, or WEBP.")

uploaded_resume = st.file_uploader(
    "Choose a resume from your device",
    type=None,
    key="resume_uploader",
    accept_multiple_files=False,
    help="Tap Browse or Files on mobile and select a PDF, DOCX, TXT, JPG, JPEG, PNG, or WEBP file. Other file types will be rejected.",
)

process_uploaded_resume(uploaded_resume)

if st.session_state.get("resume_text"):
    st.success(f"Resume loaded: {st.session_state.resume_name}")
elif st.session_state.get("resume_error"):
    st.error(f"Could not read the resume: {st.session_state.resume_error}")
elif uploaded_resume is not None:
    st.warning("No readable text was found in that file.")

if st.session_state.get("resume_text"):
    st.caption("Resume loaded and ready for questions.")

# ============================================================
# LOAD KNOWLEDGE BASE
# ============================================================

chunks, vectorizer, vectors = load_knowledge_base()


if not chunks:

    st.error(
        f"Knowledge base file not found: {KNOWLEDGE_FILE}"
    )

    st.stop()


# ============================================================
# CHECK GROQ CONFIGURATION
# ============================================================

try:

    get_groq_client()

except ValueError:

    st.error(
        "GROQ_API_KEY is not configured. Add it in Streamlit Cloud "
        "under Settings > Secrets, then reboot the app."
    )

    st.stop()


# ============================================================
# ATS RESUME GENERATOR
# ============================================================

st.subheader("Generate an ATS-friendly resume")
st.caption("Use an uploaded resume, add your details below, or combine both.")

with st.form("ats_resume_form"):
    resume_details = st.text_area(
        "Your details and target role",
        placeholder=(
            "Target role and location\n"
            "Employment history, achievements, education, projects, skills, and certifications"
        ),
        height=180,
        help="Add details that are missing from your uploaded resume. The generator will not invent information.",
    )
    generate_resume_submitted = st.form_submit_button(
        "Generate ATS resume",
        use_container_width=True,
    )

if generate_resume_submitted:
    existing_resume = st.session_state.get("resume_text", "")
    if not existing_resume and not resume_details.strip():
        st.warning("Upload a resume or enter your details before generating.")
    else:
        with st.spinner("Creating your ATS-friendly resume..."):
            try:
                st.session_state.generated_ats_resume = generate_ats_resume(
                    existing_resume,
                    resume_details,
                )
            except Exception as error:
                st.error(f"Could not generate the resume: {error}")

generated_ats_resume = st.session_state.get("generated_ats_resume", "")
if generated_ats_resume:
    st.success("ATS-friendly resume generated. Review every detail before submitting applications.")
    st.text_area(
        "Generated resume",
        value=generated_ats_resume,
        height=520,
        key="generated_ats_resume_preview",
    )
    download_columns = st.columns(2)
    with download_columns[0]:
        st.download_button(
            "Download TXT",
            data=generated_ats_resume,
            file_name="ats_resume.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with download_columns[1]:
        st.download_button(
            "Download DOCX",
            data=create_resume_docx(generated_ats_resume),
            file_name="ats_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )


# ============================================================
# CHAT HISTORY
# ============================================================

if "messages" not in st.session_state:

    st.session_state.messages = []

if not st.session_state.messages:
    st.markdown("""
    <div class="welcome">
        <div class="welcome-kicker">Your next move starts here</div>
        <h2>Turn career uncertainty<br>into a clear plan.</h2>
        <p>Ask a question, upload your resume, or pick a starting point. I will help you think it through step by step.</p>
    </div>
    """, unsafe_allow_html=True)

    st.subheader("Start with a question")
    starter_questions = [
        "What career path fits my current skills?",
        "How can I improve my resume?",
        "Build me a 30-day learning roadmap",
        "Give me 5 interview questions to practice",
    ]
    prompt_columns = st.columns(2)
    for index, starter_question in enumerate(starter_questions):
        with prompt_columns[index % 2]:
            if st.button(starter_question, key=f"starter_{index}", use_container_width=True):
                st.session_state.pending_question = starter_question
                st.rerun()

    if st.session_state.get("resume_text"):
        st.caption("Resume loaded. Try asking: “What are the strongest parts of my resume?”")


for message in st.session_state.messages:

    with st.chat_message(message["role"]):

        st.markdown(message["content"])


# ============================================================
# USER INPUT
# ============================================================

pending_question = st.session_state.pop("pending_question", None)
typed_question = st.chat_input("Ask about your career or resume...")
user_question = typed_question or pending_question


if user_question:

    # Display user message

    st.session_state.messages.append(
        {
            "role": "user",
            "content": user_question
        }
    )

    with st.chat_message("user"):

        st.markdown(user_question)


    # Generate AI response

    with st.chat_message("assistant"):

        with st.spinner(
            "AI Career Assistant is thinking..."
        ):

            try:

                knowledge_context = search_knowledge(

                    user_question,

                    chunks,

                    vectorizer,

                    vectors

                )

                resume_context = st.session_state.get("resume_text", "")
                if resume_context:
                    knowledge_context = (
                        f"General career guidance:\n{knowledge_context}\n\n"
                        "Uploaded resume:\n"
                        f"{resume_context[:12000]}"
                    )

                web_results = []
                if live_search_enabled:
                    web_results, web_context = search_web(
                        user_question,
                        resume_context,
                    )
                    if web_context:
                        knowledge_context = (
                            f"{knowledge_context}\n\n"
                            "Current public web search results (verify before relying on them):\n"
                            f"{web_context}"
                        )


                response = get_ai_response(

                    user_question,

                    knowledge_context,

                    st.session_state.messages[:-1]

                )

                answer = st.write_stream(
                    get_response_text(chunk)
                    for chunk in response
                )


                st.session_state.messages.append(

                    {
                        "role": "assistant",
                        "content": answer
                    }

                )

                if web_results:
                    with st.expander("Current web sources", expanded=False):
                        st.caption("Search results retrieved now. Check each source for the latest details.")
                        for result in web_results:
                            st.link_button(
                                result["title"][:100],
                                result["url"],
                                use_container_width=True,
                            )


            except Exception as error:

                error_message = f"""
Error while generating the response:

{error}

Please check that your GROQ_API_KEY is configured correctly
and that the Groq service is available.
"""

                st.error(error_message)